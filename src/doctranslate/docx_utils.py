from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

_PROTECTED_FIELD_PREFIXES = ("PAGEREF", "PAGE", "NUMPAGES", "SEQ")


def _is_inside_table(el) -> bool:
    parent = el.getparent()
    while parent is not None:
        if parent.tag == qn("w:tbl"):
            return True
        if parent.tag == qn("w:body"):
            return False
        parent = parent.getparent()
    return False


def iter_body_paragraphs(doc):
    """Yield every top-level (non-table) paragraph in document order, including
    ones nested inside a w:sdt content control.

    Word wraps an inserted Table of Contents field in a w:sdt ("Table of
    Contents" content control) by default. python-docx's own `doc.paragraphs`
    only returns direct children of w:body, so it skips a ToC entirely rather
    than just its hyperlink-wrapped run text. Table paragraphs are excluded
    here since callers walk `doc.tables` separately.
    """
    for el in doc.element.body.iter(qn("w:p")):
        if not _is_inside_table(el):
            yield Paragraph(el, doc)


def _is_protected_field(instr_text: str) -> bool:
    return instr_text.strip().upper().startswith(_PROTECTED_FIELD_PREFIXES)


def build_field_protection_map(doc) -> dict:
    """Map each w:r element (by its XPath, since lxml doesn't guarantee stable
    Python object identity for the same node across separate tree traversals) in
    the document body to whether it's a computed page-number-style field result
    (PAGEREF/PAGE/NUMPAGES/SEQ) that must be left untouched rather than
    translated or cleared.

    A field's begin/separate/end markers can span many paragraphs — e.g. a
    multi-entry Word Table of Contents is one continuous field wrapping every
    heading link between its begin and end — so state is tracked across the
    whole document body rather than reset per paragraph. A field is classified
    by its instrText once its own `separate` marker is reached; only its own
    result content (until its matching `end`) is protected, so a TOC field's
    own rendered entries stay translatable while a nested PAGEREF field's
    cached page number does not.
    """
    tree = doc.element.getroottree()
    protection: dict[str, bool] = {}
    stack: list[bool] = []
    pending_instr: list[str] = []
    for r in doc.element.body.iter(qn("w:r")):
        fld = r.find(qn("w:fldChar"))
        if fld is not None:
            fld_type = fld.get(qn("w:fldCharType"))
            if fld_type == "begin":
                stack.append(False)
                pending_instr = []
            elif fld_type == "separate" and stack:
                stack[-1] = _is_protected_field("".join(pending_instr))
            elif fld_type == "end" and stack:
                stack.pop()
            continue
        instr_el = r.find(qn("w:instrText"))
        if instr_el is not None:
            pending_instr.append(instr_el.text or "")
            continue
        protection[tree.getpath(r)] = bool(stack) and stack[-1]
    return protection


def iter_paragraph_runs(paragraph, protection: dict):
    """Yield (Run, is_protected) for every real-text run in document order,
    including ones nested inside w:hyperlink (python-docx's own `paragraph.runs`
    only sees direct children of w:p and silently skips these — e.g. every entry
    in a real Word Table of Contents, which links to its heading).
    """
    tree = paragraph._p.getroottree()
    for r in paragraph._p.iter(qn("w:r")):
        if r.find(qn("w:fldChar")) is not None or r.find(qn("w:instrText")) is not None:
            continue
        yield Run(r, paragraph), protection.get(tree.getpath(r), False)


def paragraph_text(paragraph, protection: dict) -> str:
    """Like `paragraph.text`, but also includes hyperlink-wrapped runs (e.g. Table
    of Contents entries) and excludes protected field results (e.g. TOC page
    numbers) that aren't real static content.
    """
    return "".join(run.text for run, protected in iter_paragraph_runs(paragraph, protection) if not protected)


def replace_paragraph_text(paragraph, text: str, protection: dict) -> None:
    """Write `text` into a paragraph's dominant real-text run (including inside
    hyperlinks), clearing other non-blank text runs but leaving tab characters
    and protected field results (e.g. TOC page numbers) untouched.
    """
    runs = [run for run, protected in iter_paragraph_runs(paragraph, protection) if not protected]
    if not runs:
        paragraph.add_run().text = text
        return
    non_blank = [r for r in runs if r.text.strip()] or runs
    dominant = max(non_blank, key=lambda r: len(r.text))
    dominant.text = text
    for run in runs:
        if run is not dominant and run.text.strip():
            run.text = ""

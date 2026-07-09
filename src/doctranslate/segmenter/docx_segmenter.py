from docx import Document

from ..docx_utils import build_field_protection_map, iter_body_paragraphs, paragraph_text
from ..models import ExtractionResult, Segment


def extract(path: str) -> ExtractionResult:
    doc = Document(path)
    protection = build_field_protection_map(doc)
    segments: list[Segment] = []
    refs: dict[str, object] = {}
    order = 0

    def add_segment(paragraph, group_key: str) -> None:
        nonlocal order
        text = paragraph_text(paragraph, protection)
        if not text.strip():
            return
        seg = Segment(doc_type="docx", source_text=text, group_key=group_key, order_hint=order)
        order += 1
        segments.append(seg)
        refs[seg.id] = paragraph

    for para in iter_body_paragraphs(doc):
        add_segment(para, "body")

    for t_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    add_segment(para, f"table{t_idx}")

    return ExtractionResult(doc_handle=doc, segments=segments, refs=refs)

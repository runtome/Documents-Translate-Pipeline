from docx import Document

from doctranslate.segmenter import docx_segmenter


def test_extract_assigns_stable_uuids_not_derived_from_position(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First")
    doc.add_paragraph("Second")
    doc.save(src)

    result = docx_segmenter.extract(str(src))

    ids = [seg.id for seg in result.segments]
    assert len(ids) == len(set(ids))
    for seg_id in ids:
        assert seg_id in result.refs


def test_extract_skips_empty_paragraphs(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Real text")
    doc.add_paragraph("")
    doc.save(src)

    result = docx_segmenter.extract(str(src))

    assert len(result.segments) == 1
    assert result.segments[0].source_text == "Real text"


def test_extract_includes_table_cells_with_distinct_group_key(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Body text")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell text"
    doc.save(src)

    result = docx_segmenter.extract(str(src))

    group_keys = {seg.group_key for seg in result.segments}
    assert "body" in group_keys
    assert "table0" in group_keys

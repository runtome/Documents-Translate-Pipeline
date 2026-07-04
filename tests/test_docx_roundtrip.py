from docx import Document

from doctranslate.pipeline import run_pipeline


def test_docx_roundtrip(tmp_path, fake_llm_client):
    src = tmp_path / "sample.docx"
    original_texts = ["Hello world", "Second paragraph"]

    doc = Document()
    for text in original_texts:
        doc.add_paragraph(text)
    doc.save(src)

    out_path = run_pipeline(str(src), "en", "th", fake_llm_client)

    assert out_path.exists()
    assert out_path != src

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]

    assert len(out_texts) == len(original_texts)
    for original, translated in zip(original_texts, out_texts):
        assert translated == f"[TR]{original}"
        assert translated != original


def test_docx_roundtrip_preserves_table_structure(tmp_path, fake_llm_client):
    src = tmp_path / "with_table.docx"
    doc = Document()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "Beta"
    doc.save(src)

    out_path = run_pipeline(str(src), "en", "ja", fake_llm_client)

    out_doc = Document(str(out_path))
    assert len(out_doc.tables) == 1
    out_table = out_doc.tables[0]
    assert len(out_table.rows) == 2
    assert len(out_table.columns) == 2
    assert out_table.cell(0, 0).text == "[TR]Name"
    assert out_table.cell(1, 1).text == "[TR]Beta"

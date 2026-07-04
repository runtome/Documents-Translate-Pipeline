import re

from docx import Document

from doctranslate.llm.base import LLMClient, LLMConfig
from doctranslate.pipeline import run_pipeline

_SEG_PATTERN = re.compile(r'<SEG id="([^"]+)">(.*?)</SEG>', re.DOTALL)


class CapturingClient(LLMClient):
    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))
        self.system_prompts: list[str] = []

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        self.system_prompts.append(system_prompt)
        parts = [f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in _SEG_PATTERN.findall(user_prompt)]
        return "\n".join(parts)


def test_glossary_terms_reach_the_system_prompt(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Alice works at the company")
    doc.save(src)

    client = CapturingClient()
    run_pipeline(str(src), "en", "th", client, glossary={"Alice": "Alice-glossary-term"})

    assert any("Alice-glossary-term" in prompt for prompt in client.system_prompts)


def test_glossary_omitted_when_no_terms_match(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Nothing relevant here")
    doc.save(src)

    client = CapturingClient()
    run_pipeline(str(src), "en", "th", client, glossary={"Alice": "Alice-glossary-term"})

    assert not any("Alice-glossary-term" in prompt for prompt in client.system_prompts)

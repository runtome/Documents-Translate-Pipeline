from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from lxml import etree

from doctranslate.pipeline import run_pipeline

# A real Word-inserted Table of Contents: the whole field lives inside a w:sdt
# content control, and each entry's visible heading text sits inside a
# w:hyperlink (linking to the heading) alongside a nested PAGEREF field whose
# cached page number is a *computed* value, not literal document content.
_TOC_SDT_XML = f"""
<w:sdt {nsdecls("w")}>
  <w:sdtPr><w:alias w:val="Table of Contents"/><w:id w:val="1"/></w:sdtPr>
  <w:sdtContent>
    <w:p>
      <w:pPr><w:pStyle w:val="TOC1"/></w:pPr>
      <w:r><w:fldChar w:fldCharType="begin"/></w:r>
      <w:r><w:instrText>TOC \\h \\o "1-3"</w:instrText></w:r>
      <w:r><w:fldChar w:fldCharType="separate"/></w:r>
      <w:hyperlink w:anchor="_Toc1" w:history="1">
        <w:r><w:t>Introduction</w:t></w:r>
        <w:r><w:tab/></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc1 \\h </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>2</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
      </w:hyperlink>
    </w:p>
    <w:p><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>
  </w:sdtContent>
</w:sdt>
"""


def test_toc_heading_text_is_translated_but_page_number_is_left_alone(tmp_path, fake_llm_client):
    src = tmp_path / "with_toc.docx"
    doc = Document()
    intro = doc.add_paragraph("Intro paragraph")
    doc.add_paragraph("Outro paragraph")
    intro._p.addnext(parse_xml(_TOC_SDT_XML))
    doc.save(src)

    out_path = run_pipeline(str(src), "en", "th", fake_llm_client)

    out_doc = Document(str(out_path))
    xml = etree.tostring(out_doc.element, encoding="unicode")

    # the TOC entry's heading label, nested inside w:hyperlink and w:sdt, must
    # have been extracted and translated despite python-docx's own
    # `paragraph.runs`/`doc.paragraphs` not seeing either wrapper
    assert "<w:t>[TR]Introduction</w:t>" in xml

    # the PAGEREF field's cached page number is a computed value, not literal
    # text -- it must survive untouched (not translated, not cleared)
    assert "<w:t>2</w:t>" in xml
    assert "[TR]2" not in xml

    # ordinary body paragraphs on either side of the sdt still translate normally
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == ["[TR]Intro paragraph", "[TR]Outro paragraph"]


def test_paragraph_text_excludes_pageref_result_but_includes_hyperlink_label():
    from doctranslate.docx_utils import build_field_protection_map, iter_body_paragraphs, paragraph_text

    doc = Document()
    p = doc.add_paragraph("placeholder")
    p._p.addnext(parse_xml(_TOC_SDT_XML))

    protection = build_field_protection_map(doc)
    toc_paragraph = next(
        para for para in iter_body_paragraphs(doc) if para._p.find(qn("w:hyperlink")) is not None
    )
    assert paragraph_text(toc_paragraph, protection) == "Introduction\t"

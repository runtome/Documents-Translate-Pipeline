import re

from docx import Document

from doctranslate.llm.base import LLMClient, LLMConfig
from doctranslate.pipeline import run_pipeline

_SEG_PATTERN = re.compile(r'<SEG id="([^"]+)">(.*?)</SEG>', re.DOTALL)
_UUID_HEX_PATTERN = re.compile(r"\b[0-9a-f]{32}\b")


class CapturingClient(LLMClient):
    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))
        self.user_prompts: list[str] = []

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        self.user_prompts.append(user_prompt)
        parts = [f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in _SEG_PATTERN.findall(user_prompt)]
        return "\n".join(parts)


def test_prompt_uses_short_aliases_not_raw_uuids(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(5):
        doc.add_paragraph(f"Paragraph number {i}")
    doc.save(src)

    client = CapturingClient()
    out_path = run_pipeline(str(src), "en", "th", client)

    assert client.user_prompts, "expected at least one chunk to be sent"
    for prompt in client.user_prompts:
        assert not _UUID_HEX_PATTERN.search(prompt), "raw 32-char hex ids should never reach the LLM prompt"

    first_prompt_ids = [seg_id for seg_id, _ in _SEG_PATTERN.findall(client.user_prompts[0])]
    assert first_prompt_ids == [str(i) for i in range(1, len(first_prompt_ids) + 1)]

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    for i, text in enumerate(out_texts):
        assert text == f"[TR]Paragraph number {i}"

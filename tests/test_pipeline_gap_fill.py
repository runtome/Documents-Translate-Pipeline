import re

import pytest
from docx import Document

from doctranslate.exceptions import ChunkProcessingError
from doctranslate.llm.base import LLMClient, LLMConfig
from doctranslate.pipeline import run_pipeline

_SEG_PATTERN = re.compile(r'<SEG id="([^"]+)">(.*?)</SEG>', re.DOTALL)


class DropsOneInBatchClient(LLMClient):
    """Always drops alias "3" when asked to translate more than one segment at once,
    but translates correctly when given a single segment - simulating a small local
    model that reliably fumbles exactly one tag out of a big batch, no matter how many
    times the whole batch is retried, yet handles a lone segment fine.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))
        self.requests: list[list[str]] = []

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        self.requests.append([seg_id for seg_id, _ in matches])
        parts = []
        for seg_id, text in matches:
            if len(matches) > 1 and seg_id == "3":
                continue
            parts.append(f'<SEG id="{seg_id}">[TR]{text}</SEG>')
        return "\n".join(parts)


def test_gap_fill_recovers_single_dropped_segment_without_losing_the_rest(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(5):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = DropsOneInBatchClient()
    out_path = run_pipeline(str(src), "en", "th", client, max_retries=3)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == [f"[TR]Paragraph {i}" for i in range(5)]

    # first request covers all 5, then retries the full batch, then a final
    # single-segment gap-fill request recovers the one that kept getting dropped
    assert len(client.requests) >= 2
    assert client.requests[-1] == ["1"]


class AlwaysDropsOneOfTwoClient(LLMClient):
    """Even the gap-fill sub-chunk still drops one id - the missing segment should be
    left untranslated (not silently invented), and on-error=skip should keep the rest.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        parts = [f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in matches[:-1]]
        return "\n".join(parts)


def test_unrecoverable_segment_is_left_untranslated_when_on_error_is_skip(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(src)

    client = AlwaysDropsOneOfTwoClient()
    out_path = run_pipeline(str(src), "en", "th", client, on_error="skip", max_retries=1)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == ["[TR]First paragraph", "Second paragraph"]


class UnclosedTrailingTagClient(LLMClient):
    """Simulates a model that forgets to close the very last <SEG> tag in a
    batch response (e.g. it hit a stop/token limit mid-generation). This is
    flagged malformed (raw "<SEG" count != parsed pairs), but since nothing
    follows the unclosed tag, exactly one id is genuinely absent from the
    parsed translations and nothing else is corrupted - it should still be
    gap-filled. Closes tags correctly when given a single segment.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        parts = []
        for i, (seg_id, text) in enumerate(matches):
            if len(matches) > 1 and i == len(matches) - 1:
                parts.append(f'<SEG id="{seg_id}">[TR]{text}')  # deliberately unclosed
            else:
                parts.append(f'<SEG id="{seg_id}">[TR]{text}</SEG>')
        return "\n".join(parts)


def test_gap_fill_recovers_from_a_malformed_response_with_an_unclosed_trailing_tag(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(5):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = UnclosedTrailingTagClient()
    out_path = run_pipeline(str(src), "en", "th", client, max_retries=1)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == [f"[TR]Paragraph {i}" for i in range(5)]


class CorruptedMergeClient(LLMClient):
    """Leaves a *non-trailing* <SEG> tag unclosed whenever given 2+ segments, so
    the parser's non-greedy match swallows the next tag's opener into this
    segment's content before finding a real closing tag - producing a
    corrupted (not merely missing) value every time it's given 2+ segments,
    but translating cleanly for a single segment. Since the retry set here
    never shrinks below the full remaining batch (quarantining "corrupts"
    just relabels the same ids as "missing"), plain gap-fill alone can't
    converge - it takes the split-in-half fallback bottoming out at
    single-segment requests to fully recover this.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        if len(matches) == 1:
            seg_id, text = matches[0]
            return f'<SEG id="{seg_id}">[TR]{text}</SEG>'
        parts = []
        for i, (seg_id, text) in enumerate(matches):
            if i == 0:
                parts.append(f'<SEG id="{seg_id}">[TR]{text}')  # unclosed, not last
            else:
                parts.append(f'<SEG id="{seg_id}">[TR]{text}</SEG>')
        return "\n".join(parts)


def test_split_in_half_fallback_recovers_a_never_shrinking_corruption(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(3):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = CorruptedMergeClient()
    out_path = run_pipeline(str(src), "en", "th", client, max_retries=1)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == [f"[TR]Paragraph {i}" for i in range(3)]


class AlwaysFailsClient(LLMClient):
    """Never produces a single valid <SEG> pair, regardless of batch size -
    even a lone segment always comes back empty. This is the true
    nothing-more-can-be-tried case that must still raise rather than split
    or retry forever.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        return "no valid tags in this response at all"


def test_client_that_always_fails_even_alone_is_not_retried_forever(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(3):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = AlwaysFailsClient()

    with pytest.raises(ChunkProcessingError):
        run_pipeline(str(src), "en", "th", client, max_retries=1)


class FailsAboveThresholdClient(LLMClient):
    """Returns a response with no valid tags at all whenever asked to
    translate more than `threshold` segments at once; translates correctly
    otherwise - simulating a small/weak model (or a model given too large a
    chunk for its context/output window) that can't handle a big batch at
    all, as opposed to just fumbling one or two ids within it.
    """

    def __init__(self, threshold: int = 2):
        super().__init__(LLMConfig(provider="fake", model="fake"))
        self.threshold = threshold
        self.batch_sizes: list[int] = []

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        self.batch_sizes.append(len(matches))
        if len(matches) > self.threshold:
            return "sorry, I can't do that many at once"
        return "\n".join(f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in matches)


def test_total_batch_failure_is_recovered_by_splitting_in_half(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(5):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = FailsAboveThresholdClient(threshold=2)
    out_path = run_pipeline(str(src), "en", "th", client, max_retries=1)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == [f"[TR]Paragraph {i}" for i in range(5)]

    # the initial full-size attempt fails outright, but every batch actually
    # accepted (i.e. that produced output) is at or below the client's threshold
    assert client.batch_sizes[0] == 5
    accepted_sizes = [size for size in client.batch_sizes if size <= client.threshold]
    assert accepted_sizes and max(accepted_sizes) <= 2


class CorruptedMergeThenRecoversClient(LLMClient):
    """Only corrupts (merges the first tag into the second's close) for
    batches of 3+ segments; translates cleanly for smaller batches -
    simulating a small model that trips over its own tag structure on longer
    runs but succeeds once gap-fill narrows the retry down to just the
    affected ids (a realistic, non-adversarial version of the corrupted-merge
    failure mode).
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        if len(matches) < 3:
            return "\n".join(f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in matches)
        parts = []
        for i, (seg_id, text) in enumerate(matches):
            if i == 0:
                parts.append(f'<SEG id="{seg_id}">[TR]{text}')  # unclosed, not last
            else:
                parts.append(f'<SEG id="{seg_id}">[TR]{text}</SEG>')
        return "\n".join(parts)


def test_gap_fill_recovers_by_quarantining_a_corrupted_value_and_retrying_it_too(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(4):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = CorruptedMergeThenRecoversClient()
    out_path = run_pipeline(str(src), "en", "th", client, max_retries=1)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == [f"[TR]Paragraph {i}" for i in range(4)]


class InventsOneExtraSegmentClient(LLMClient):
    """Always returns every requested segment correctly, plus one bonus,
    unrequested id past the end of the batch - simulating a small model that
    hallucinates one extra numbered entry that doesn't exist in the source.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))
        self.call_count = 0

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        self.call_count += 1
        matches = _SEG_PATTERN.findall(user_prompt)
        parts = [f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in matches]
        bonus_id = str(len(matches) + 1)
        parts.append(f'<SEG id="{bonus_id}">[TR]hallucinated content</SEG>')
        return "\n".join(parts)


def test_harmless_extra_id_is_discarded_without_retrying_or_aborting(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(3):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = InventsOneExtraSegmentClient()
    out_path = run_pipeline(str(src), "en", "th", client)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == [f"[TR]Paragraph {i}" for i in range(3)]
    assert client.call_count == 1


class HallucinatesCorruptedPlaceholderIdClient(LLMClient):
    """Translates every real segment correctly, but also appends a bogus
    non-numeric id (e.g. a model echoing a literal "..." placeholder from the
    prompt's example format) whose own content happens to look corrupted -
    this must not crash `sorted(..., key=int)` and must not be treated as a
    reason to distrust the otherwise-complete response.
    """

    def __init__(self):
        super().__init__(LLMConfig(provider="fake", model="fake"))

    def translate_chunk(self, system_prompt: str, user_prompt: str) -> str:
        matches = _SEG_PATTERN.findall(user_prompt)
        parts = [f'<SEG id="{seg_id}">[TR]{text}</SEG>' for seg_id, text in matches]
        parts.append('<SEG id="...">junk <SEG more junk</SEG>')
        return "\n".join(parts)


def test_non_numeric_corrupted_placeholder_id_does_not_crash_or_fail(tmp_path):
    src = tmp_path / "sample.docx"
    doc = Document()
    for i in range(2):
        doc.add_paragraph(f"Paragraph {i}")
    doc.save(src)

    client = HallucinatesCorruptedPlaceholderIdClient()
    out_path = run_pipeline(str(src), "en", "th", client)

    out_doc = Document(str(out_path))
    out_texts = [p.text for p in out_doc.paragraphs if p.text.strip()]
    assert out_texts == ["[TR]Paragraph 0", "[TR]Paragraph 1"]


def test_persistently_failing_single_segment_logs_its_source_text(tmp_path, caplog):
    src = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("A stubborn paragraph that the model refuses to translate")
    doc.save(src)

    client = AlwaysFailsClient()

    with caplog.at_level("WARNING"):
        with pytest.raises(ChunkProcessingError):
            run_pipeline(str(src), "en", "th", client, on_error="abort", max_retries=1)

    assert "A stubborn paragraph that the model refuses to translate" in caplog.text
